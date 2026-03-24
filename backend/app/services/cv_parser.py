"""
services/cv_parser.py
======================
PDF and DOCX text extraction for uploaded CV files.
"""
from __future__ import annotations

import io

from app.config import settings


def validate_file_size(file_bytes: bytes, max_mb: int | None = None) -> bool:
    """Return True if file size is within the allowed limit."""
    limit = (max_mb or settings.MAX_UPLOAD_SIZE_MB) * 1024 * 1024
    return len(file_bytes) <= limit


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF file using pdfplumber.
    Falls back to PyPDF2 if pdfplumber fails.
    """
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        return "\n".join(pages_text).strip()
    except Exception:
        # Fallback: PyPDF2
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file using python-docx.
    Extracts from paragraphs and table cells.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts).strip()


def parse_cv_file(filename: str, file_bytes: bytes) -> str:
    """
    Route-aware CV parser: choose PDF or DOCX extractor based on file extension.

    Raises
    ------
    ValueError
        If the file extension is not .pdf or .docx.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Only PDF and DOCX are accepted.")
